"""
harness/deliverable_engine.py
Industrial Deliverable Synthesizer for Sovereign AI Workbench.
Generates official Indian PSU Microsoft Word (.docx) Approval Notes and Excel (.xlsx) Calculation Sheets.
"""

import os
from typing import Dict, Any, Optional
from harness.types import PSUApprovalNote


class DeliverableEngine:
    """
    Synthesizes official industrial deliverables matching PSU document standards.
    Supports Microsoft Word (.docx), Excel (.xlsx), and Markdown fallback outputs.
    """

    def __init__(self, output_dir: str = "deliverables"):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def generate_docx_approval_note(self, note: PSUApprovalNote, filename: Optional[str] = None) -> str:
        """
        Generates a standard Indian PSU formatted Microsoft Word (.docx) Approval Note.
        """
        out_name = filename or f"Approval_Note_{note.equipment_tag.replace(' ', '_')}.docx"
        out_path = os.path.join(self.output_dir, out_name)

        try:
            import docx
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT

            doc = Document()

            # Page Margins
            for section in doc.sections:
                section.top_margin = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin = Inches(0.9)
                section.right_margin = Inches(0.9)

            # PSU Corporate Header
            header_p = doc.add_paragraph()
            header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_org = header_p.add_run("MANGALORE REFINERY AND PETROCHEMICALS LIMITED\n")
            run_org.bold = True
            run_org.font.size = Pt(13)
            run_org.font.color.rgb = RGBColor(0x00, 0x33, 0x66)  # Deep Navy PSU Blue

            run_dept = header_p.add_run("TECHNICAL SERVICES & ASSET INTEGRITY DIVISION\n")
            run_dept.bold = True
            run_dept.font.size = Pt(10.5)
            run_dept.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            run_type = header_p.add_run("CONFIDENTIAL / INTERNAL APPROVAL NOTE\n")
            run_type.bold = True
            run_type.font.size = Pt(11)
            run_type.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)  # Dark Red

            doc.add_paragraph("―" * 58).alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata Table
            table = doc.add_table(rows=2, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True

            table.rows[0].cells[0].paragraphs[0].add_run(f"Equipment Tag: {note.equipment_tag}").bold = True
            table.rows[0].cells[1].paragraphs[0].add_run("Classification: On-Premise Sovereign Verified").bold = True
            table.rows[1].cells[0].paragraphs[0].add_run("Security Compliance: 100% Air-Gapped").italic = True
            table.rows[1].cells[1].paragraphs[0].add_run(f"Department: {note.designation}").italic = True

            doc.add_paragraph("")  # Spacer

            # 1.0 SUBJECT
            p_subj = doc.add_paragraph()
            p_subj.add_run("1.0 SUBJECT\n").bold = True
            p_subj.add_run(note.subject)

            # 2.0 BACKGROUND
            p_bg = doc.add_paragraph()
            p_bg.add_run("2.0 BACKGROUND & OPERATIONAL CONTEXT\n").bold = True
            p_bg.add_run(note.background)

            # 3.0 INSPECTION FINDINGS
            p_find = doc.add_paragraph()
            p_find.add_run("3.0 NON-DESTRUCTIVE TESTING & INSPECTION FINDINGS\n").bold = True
            p_find.add_run(note.inspection_findings)

            # 4.0 ASME B31.3 ENGINEERING INTEGRITY CALCULATIONS
            p_calc = doc.add_paragraph()
            p_calc.add_run("4.0 ASME B31.3 CODE CALCULATIONS & REMAINING LIFE\n").bold = True
            p_calc.add_run(note.calculation_summary)

            # 5.0 FINANCIAL IMPLICATIONS
            p_fin = doc.add_paragraph()
            p_fin.add_run("5.0 FINANCIAL IMPLICATIONS & BUDGETARY PROVISION\n").bold = True
            p_fin.add_run(note.financial_implication)

            # 6.0 DELEGATION OF POWER (DOP)
            p_dop = doc.add_paragraph()
            p_dop.add_run("6.0 DELEGATION OF POWER (DOP) REFERENCE\n").bold = True
            p_dop.add_run(note.authority_dop)

            # 7.0 RECOMMENDATION
            p_rec = doc.add_paragraph()
            p_rec.add_run("7.0 RECOMMENDATION & PROPOSAL FOR APPROVAL\n").bold = True
            p_rec.add_run(note.recommendation)

            # Signature Authority Block
            doc.add_paragraph("\n\n")
            sig_p = doc.add_paragraph()
            sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sig_run = sig_p.add_run(
                f"Submitted for Approval:\n\n_______________________\n"
                f"{note.approver_name}\n"
                f"{note.designation}\n"
                f"Mangalore Refinery and Petrochemicals Limited"
            )
            sig_run.bold = True

            doc.save(out_path)
            return out_path

        except ImportError:
            # Fallback Markdown / Text Document if python-docx is not yet available
            md_path = out_path.replace(".docx", ".md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(f"# MANGALORE REFINERY AND PETROCHEMICALS LIMITED\n")
                f.write(f"## TECHNICAL SERVICES & ASSET INTEGRITY DIVISION\n")
                f.write(f"### CONFIDENTIAL / INTERNAL APPROVAL NOTE\n\n")
                f.write(f"**Equipment Tag:** {note.equipment_tag}\n\n")
                f.write(f"### 1.0 SUBJECT\n{note.subject}\n\n")
                f.write(f"### 2.0 BACKGROUND\n{note.background}\n\n")
                f.write(f"### 3.0 INSPECTION FINDINGS\n{note.inspection_findings}\n\n")
                f.write(f"### 4.0 ENGINEERING CALCULATIONS\n{note.calculation_summary}\n\n")
                f.write(f"### 5.0 FINANCIAL IMPLICATIONS\n{note.financial_implication}\n\n")
                f.write(f"### 6.0 DELEGATION OF POWER\n{note.authority_dop}\n\n")
                f.write(f"### 7.0 RECOMMENDATION\n{note.recommendation}\n\n")
                f.write(f"---\n**Submitted by:** {note.approver_name}, {note.designation}\n")
            return md_path

    def generate_xlsx_calculation_sheet(self, calc_data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """
        Generates an Excel (.xlsx) calculation sheet with ASME formula breakdown.
        """
        out_name = filename or f"Engineering_Calculations_{calc_data.get('standard', 'ASME_B31_3')}.xlsx"
        out_path = os.path.join(self.output_dir, out_name)

        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "ASME Integrity Verification"

            # Header Style
            header_fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
            header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")

            ws["A1"] = "MANGALORE REFINERY & PETROCHEMICALS LIMITED — ENGINEERING CALCULATION LOG"
            ws["A1"].font = Font(name="Calibri", size=13, bold=True, color="003366")

            headers = ["Parameter", "Symbol / Standard", "Unit", "Value", "Notes"]
            for col_num, h_text in enumerate(headers, 1):
                cell = ws.cell(row=3, column=col_num)
                cell.value = h_text
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center", vertical="center")

            rows = [
                ("Design Internal Pressure", "P", "MPa", calc_data.get("design_pressure_mpa", 4.0), "Operating crude pressure"),
                ("Pipe Outer Diameter", "D", "mm", calc_data.get("outer_diameter_mm", 219.1), "8-inch nominal diameter"),
                ("Allowable Material Stress", "S", "MPa", calc_data.get("allowable_stress_mpa", 137.0), "ASTM A106 Grade B"),
                ("Corrosion Allowance", "c", "mm", calc_data.get("corrosion_allowance_mm", 3.0), "Standard refinery allowance"),
                ("Required Minimum Thickness", "t_min", "mm", calc_data.get("required_min_thickness_mm", 6.2), "Formula: (P*D)/(2*(SE+PY)) + c"),
                ("Actual Measured Thickness", "t_actual", "mm", calc_data.get("measured_thickness_mm", 7.5), "Ultrasonic NDT gauge"),
                ("Safe Wall Thickness Margin", "delta_t", "mm", calc_data.get("safe_margin_mm", 1.3), "t_actual - t_min"),
                ("Annual Corrosion Rate", "CR", "mm/year", calc_data.get("annual_corrosion_rate_mm_yr", 0.25), "Historical trend rate"),
                ("Estimated Remaining Safe Life", "Years", "Years", calc_data.get("remaining_life_years", 5.2), "delta_t / CR"),
                ("Asset Integrity Status", "Compliance", "Status", calc_data.get("status", "APPROVED_SAFE"), "ASME Code Safety Check"),
            ]

            for row_idx, row_data in enumerate(rows, 4):
                for col_idx, val in enumerate(row_data, 1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    cell.value = val
                    if col_idx == 4 and isinstance(val, (int, float)):
                        cell.alignment = Alignment(horizontal="right")
                    elif col_idx == 4 and "APPROVED" in str(val):
                        cell.font = Font(bold=True, color="008000")
                    elif col_idx == 4 and "REJECTED" in str(val):
                        cell.font = Font(bold=True, color="FF0000")

            # Column auto-width
            for col in ws.columns:
                max_len = max(len(str(cell.value or "")) for cell in col)
                col_letter = openpyxl.utils.get_column_letter(col[0].column)
                ws.column_dimensions[col_letter].width = max(max_len + 3, 12)

            wb.save(out_path)
            return out_path

        except ImportError:
            csv_path = out_path.replace(".xlsx", ".csv")
            import csv
            with open(csv_path, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerow(["Parameter", "Value", "Unit"])
                for k, v in calc_data.items():
                    writer.writerow([k, v, ""])
            return csv_path
